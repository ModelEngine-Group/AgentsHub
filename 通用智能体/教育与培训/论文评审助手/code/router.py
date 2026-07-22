"""论文评审路由 - 离线引擎 + LLM API"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from app.modules.review.schemas import PaperReviewRequest, PaperReviewResponse
from app.modules.review.offline_engine import offline_engine
import json
import os
import re
import httpx
import logging
import asyncio
from typing import Optional
from pathlib import Path
from docx import Document

from ...shared.config import get_settings

logger = logging.getLogger(__name__)
router = APIRouter(tags=["论文评审"])

# 从统一配置读取
settings = get_settings()

# 评审配置
REVIEW_CONFIGS = {
    "math_modeling": {
        "name": "数学建模论文",
        "aspects": [
            {"name": "问题分析", "max_score": 20, "weight": 20},
            {"name": "模型建立", "max_score": 30, "weight": 30},
            {"name": "求解方法", "max_score": 20, "weight": 20},
            {"name": "结果分析", "max_score": 15, "weight": 15},
            {"name": "论文写作", "max_score": 15, "weight": 15},
        ]
    },
    "academic": {
        "name": "学术论文",
        "aspects": [
            {"name": "创新性", "max_score": 25, "weight": 25},
            {"name": "方法论", "max_score": 25, "weight": 25},
            {"name": "实验与结果", "max_score": 20, "weight": 20},
            {"name": "文献综述", "max_score": 15, "weight": 15},
            {"name": "写作质量", "max_score": 15, "weight": 15},
        ]
    },
    "thesis": {
        "name": "学位论文",
        "aspects": [
            {"name": "研究意义", "max_score": 15, "weight": 15},
            {"name": "文献综述", "max_score": 20, "weight": 20},
            {"name": "研究方法", "max_score": 20, "weight": 20},
            {"name": "研究结果", "max_score": 25, "weight": 25},
            {"name": "论文规范", "max_score": 20, "weight": 20},
        ]
    },
    "course": {
        "name": "课程论文",
        "aspects": [
            {"name": "问题理解", "max_score": 20, "weight": 20},
            {"name": "内容质量", "max_score": 30, "weight": 30},
            {"name": "分析能力", "max_score": 20, "weight": 20},
            {"name": "结构组织", "max_score": 15, "weight": 15},
            {"name": "格式规范", "max_score": 15, "weight": 15},
        ]
    }
}


async def call_llm(prompt: str, max_retries: int = 2) -> str:
    """异步调用 SiliconFlow LLM API，带重试"""
    if not settings.SILICONFLOW_API_KEY:
        logger.warning("未配置 SILICONFLOW_API_KEY，使用离线评审模式")
        return generate_mock_review(prompt)

    headers = {
        "Authorization": f"Bearer {settings.SILICONFLOW_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": settings.LLM_MODEL,
        "messages": [
            {"role": "system", "content": "你是一个专业的论文评审专家，擅长对各类论文进行多维度评审。请用中文回答。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7,
        "max_tokens": 4096
    }

    for attempt in range(max_retries + 1):
        try:
            async with httpx.AsyncClient(timeout=90.0) as client:
                response = await client.post(
                    f"{settings.LLM_BASE_URL}/chat/completions",
                    headers=headers,
                    json=payload
                )
                response.raise_for_status()
                result = response.json()
                return result["choices"][0]["message"]["content"]
        except (httpx.TimeoutException, httpx.ConnectError) as e:
            logger.warning(f"LLM API 调用失败 (尝试 {attempt+1}/{max_retries+1}): {e}")
            if attempt < max_retries:
                await asyncio.sleep(2 ** attempt)
            else:
                logger.error(f"LLM API 最终失败，回退到离线模式: {e}")
                return generate_mock_review(prompt)
        except Exception as e:
            logger.error(f"LLM API 异常: {e}")
            return generate_mock_review(prompt)


def analyze_paper_content(content: str) -> dict:
    """分析论文内容，提取关键信息"""
    content = content.strip()
    word_count = len(content)
    is_empty = word_count < 50

    return {
        "is_empty": is_empty,
        "has_abstract": False if is_empty else bool(re.search(r'摘[要要]', content[:2000])),
        "has_keywords": False if is_empty else bool(re.search(r'关键[词字]', content[:2000])),
        "has_introduction": False if is_empty else bool(re.search(r'[引言绪论]|一、|1\\.', content[:3000])),
        "has_model": False if is_empty else bool(re.search(r'模型|建模|建立', content)),
        "has_algorithm": False if is_empty else bool(re.search(r'算法|求解|优化|遗传|粒子群|模拟退火', content)),
        "has_results": False if is_empty else bool(re.search(r'结果|实验|仿真|数值', content)),
        "has_figures": False if is_empty else bool(re.search(r'图\\s*\\d|表\\s*\\d|Figure|Table', content)),
        "has_references": False if is_empty else bool(re.search(r'参考文献|引用', content[-2000:])),
        "has_code": False if is_empty else bool(re.search(r'代码|程序|附录|import|def |class ', content)),
        "word_count": word_count,
        "paragraph_count": len([p for p in content.split(chr(10)) if p.strip()]),
    }


def generate_mock_review(prompt: str) -> str:
    """生成模拟评审结果（当 API 不可用时）"""
    content_match = re.search(r'论文内容\s*\n+\s*(.+?)\s*\n+\s*---', prompt, re.DOTALL)
    content = content_match.group(1) if content_match else ""
    analysis = analyze_paper_content(content)

    issues = []
    strengths = []

    if analysis["has_abstract"]:
        strengths.append("包含摘要部分")
    else:
        issues.append("缺少摘要，数模论文摘要必须包含问题、模型、算法、结果、结论五要素")

    if analysis["has_model"]:
        strengths.append("有模型建立相关内容")
    else:
        issues.append("缺少明确的模型建立过程")

    if analysis["has_algorithm"]:
        strengths.append("包含算法求解描述")
    else:
        issues.append("缺少算法求解的具体描述")

    if analysis["has_results"]:
        strengths.append("有结果分析")
    else:
        issues.append("缺少实验结果或数值分析")

    if analysis["has_figures"]:
        strengths.append("包含图表")
    else:
        issues.append("缺少图表，数模论文需要可视化展示结果")

    if analysis["has_references"]:
        strengths.append("包含参考文献")
    else:
        issues.append("缺少参考文献")

    if analysis["word_count"] < 1000:
        issues.append("论文字数偏少，建议扩充内容")

    strengths_text = "\n".join([f"- {s}" for s in strengths]) if strengths else "- 论文结构基本完整"
    issues_text = "\n".join([f"- {i}" for i in issues]) if issues else "- 整体质量较好"

    # 空内容处理
    if analysis["is_empty"]:
        return f"""# 论文评审报告

## 总体评价
**无法评审**：上传的文件未能提取到有效文本内容（{analysis['word_count']} 字）。

可能原因：
1. PDF 为扫描件/图片，无法识别文字
2. 文件格式不支持
3. 文件已损坏

## 建议
- 请上传可选中文字的 PDF 文件（非扫描件）
- 或将论文内容复制粘贴到文本框中直接评审
- 支持格式：PDF、DOCX、TXT、MD

---
*当前为离线评审模式，配置 API Key 后可获得 AI 评审。*
"""

    base_score = 30  # 起始分
    if analysis["has_abstract"]: base_score += 10
    if analysis["has_keywords"]: base_score += 3
    if analysis["has_introduction"]: base_score += 5
    if analysis["has_model"]: base_score += 12
    if analysis["has_algorithm"]: base_score += 10
    if analysis["has_results"]: base_score += 8
    if analysis["has_figures"]: base_score += 5
    if analysis["has_references"]: base_score += 5
    if analysis["has_code"]: base_score += 3
    if analysis["word_count"] > 3000: base_score += 4
    if analysis["word_count"] > 5000: base_score += 3
    if analysis["word_count"] > 8000: base_score += 2
    base_score = min(95, max(0, base_score))

    return f"""# 论文评审报告

## 总体评价
本论文共 {analysis['word_count']} 字，{analysis['paragraph_count']} 段。整体结构{'较为完整' if len(strengths) > 3 else '有待完善'}，{'模型建立和求解过程清晰' if analysis['has_model'] and analysis['has_algorithm'] else '需要加强模型建立和求解部分的论述'}。

## 各维度评审

### 1. 问题分析 ({min(20, 15 + (3 if analysis['has_introduction'] else 0))}/20分)
**优点**：
{strengths_text}

**不足**：
{issues_text}

### 2. 模型建立 ({min(30, 22 + (4 if analysis['has_model'] else 0) + (2 if analysis['has_abstract'] else 0))}/30分)
**优点**：
- {'有明确的模型描述' if analysis['has_model'] else '论文结构基本完整'}

**不足**：
- {'建议补充模型假设的合理性验证' if analysis['has_model'] else '需要详细描述模型建立过程'}

### 3. 求解方法 ({min(20, 14 + (3 if analysis['has_algorithm'] else 0) + (2 if analysis['has_code'] else 0))}/20分)
**优点**：
- {'包含算法描述' if analysis['has_algorithm'] else '有基本的求解思路'}

**不足**：
- {'建议补充算法复杂度分析' if analysis['has_algorithm'] else '需要详细描述求解算法'}

### 4. 结果分析 ({min(15, 10 + (2 if analysis['has_results'] else 0) + (2 if analysis['has_figures'] else 0))}/15分)
**优点**：
- {'有结果展示' if analysis['has_results'] else '论文结构完整'}

**不足**：
- {'建议增加灵敏度分析' if analysis['has_results'] else '需要补充实验结果'}

### 5. 论文写作 ({min(15, 10 + (2 if analysis['has_references'] else 0) + (2 if analysis['has_figures'] else 0))}/15分)
**优点**：
- {'包含参考文献' if analysis['has_references'] else '格式基本规范'}

**不足**：
- {'建议检查图表是否配有文字分析' if analysis['has_figures'] else '建议补充图表和参考文献'}

## 改进建议
1. {'确保摘要包含五要素：问题、模型、算法、结果、结论' if not analysis['has_abstract'] else '摘要可以更加精炼'}
2. {'补充模型假设和推导过程' if not analysis['has_model'] else '增加模型适用性论证'}
3. {'详细描述求解算法步骤' if not analysis['has_algorithm'] else '补充算法复杂度分析'}
4. {'增加实验结果和数值分析' if not analysis['has_results'] else '增加灵敏度分析和对比实验'}
5. {'添加图表可视化结果' if not analysis['has_figures'] else '确保每张图表都有文字分析'}

---
*注：当前为离线评审模式，配置 SiliconFlow API Key 后可获得更精准的 AI 评审。*
"""


def extract_text_from_docx(file_path: str) -> str:
    """从 DOCX 文件提取文本"""
    try:
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return "\n".join(text)
    except Exception as e:
        logger.error(f"DOCX 解析失败: {e}")
        return f"DOCX 解析失败: {str(e)}"


def _format_offline_report(result: dict) -> str:
    """格式化离线评审报告"""
    lines = ["# 论文评审报告（离线模式）\n"]

    lines.append(f"## 总体评价\n{result.get('analysis_summary', '')}\n")
    lines.append(f"**总分：{result['total_score']}/100  等级：{result['grade']}**\n")

    lines.append("## 各维度评分")
    for dim in result.get("dimensions", []):
        bar = "█" * (dim["score"] // 2) + "░" * ((dim["max_score"] - dim["score"]) // 2)
        lines.append(f"- {dim['name']}: **{dim['score']}/{dim['max_score']}** {bar}")

    if result.get("feedback"):
        lines.append("\n## 发现的问题")
        for f in result["feedback"]:
            lines.append(f"- {f}")

    if result.get("improvements"):
        lines.append("\n## 改进建议（按优先级）")
        for i, imp in enumerate(result["improvements"], 1):
            lines.append(f"{i}. **[{imp['priority']}] {imp['action']}**：{imp['detail']}")

    if result.get("recommendations"):
        lines.append("\n## 推荐模型")
        for m in result["recommendations"][:3]:
            lines.append(f"- {m['name']}（使用 {m['method']}）")

    lines.append("\n---\n*评审由离线引擎完成，配置 API Key 后可获得 AI 增强评审。*")
    return "\n".join(lines)


@router.post("/review/paper", response_model=PaperReviewResponse)
async def review_paper(request: PaperReviewRequest):
    """评审论文 - 离线引擎 + LLM"""
    paper_type = request.paper_type or "math_modeling"

    # 优先使用离线引擎（不需要API）
    if not settings.SILICONFLOW_API_KEY:
        result = offline_engine.review(request.content, paper_type)
        return PaperReviewResponse(
            paper_type=REVIEW_CONFIGS.get(paper_type, REVIEW_CONFIGS["math_modeling"])["name"],
            total_score=result["total_score"],
            grade=result["grade"],
            review_report=_format_offline_report(result),
            dimensions=result["dimensions"],
        )

    # 有API Key时用LLM评审
    config = REVIEW_CONFIGS.get(paper_type, REVIEW_CONFIGS["math_modeling"])

    aspects_desc = "\n".join([
        f"- {a['name']}（满分 {a['max_score']} 分）"
        for a in config["aspects"]
    ])

    prompt = f"""你是全国大学生数学建模竞赛的资深评委，拥有 20 年评审经验。请严格按照国赛评分标准对以下论文进行专业评审。

## 国赛评分标准（总分 100 分）

### 一、摘要（10 分）
- 是否包含五要素：问题、模型、算法、结果、结论
- 是否简明扼要（300-500 字）
- 是否突出创新点

### 二、模型建立（30 分）
- 模型假设是否合理（5 分）
- 符号说明是否规范（5 分）
- 模型推导是否严密（10 分）
- 模型是否有创新性（10 分）

### 三、求解方法（25 分）
- 算法选择是否恰当（10 分）
- 算法描述是否清晰（10 分）
- 是否有算法复杂度分析（5 分）

### 四、结果分析（20 分）
- 结果是否正确（10 分）
- 是否有灵敏度分析（5 分）
- 是否有对比实验（5 分）

### 五、论文写作（15 分）
- 逻辑是否清晰（5 分）
- 图表是否规范（5 分）
- 参考文献是否规范（5 分）

---

## 论文内容

{request.content[:8000]}

---

## 评审要求

请严格按照以下格式输出评审结果，每个维度都要给出具体分数和详细评语：

### 总体评价
（200 字以内，概述论文质量水平）

### 各维度评审

#### 1. 摘要（X/10 分）
**优点：**
- ...
**不足：**
- ...
**评语：** ...

#### 2. 模型建立（X/30 分）
**优点：**
- ...
**不足：**
- ...
**评语：** ...

#### 3. 求解方法（X/25 分）
**优点：**
- ...
**不足：**
- ...
**评语：** ...

#### 4. 结果分析（X/20 分）
**优点：**
- ...
**不足：**
- ...
**评语：** ...

#### 5. 论文写作（X/15 分）
**优点：**
- ...
**不足：**
- ...
**评语：** ...

### 改进建议（按优先级排序）
1. ...
2. ...
3. ...
4. ...
5. ...

### 国奖竞争力评估
- 当前水平：省级X等奖 / 国家级X等奖
- 主要差距：
- 提升方向：
"""

    review_content = await call_llm(prompt)

    # 空内容/无法评审处理
    if "无法评审" in review_content or not request.content.strip():
        return PaperReviewResponse(
            paper_type=config["name"],
            total_score=0,
            grade="无法评审",
            review_report=review_content,
            dimensions=[{"name": a["name"], "score": 0, "max_score": a["max_score"]} for a in config["aspects"]]
        )

    dimensions = []
    total_score = 0

    for aspect in config["aspects"]:
        score_pattern = rf'{aspect["name"]}.*?(\d+)/{aspect["max_score"]}'
        score_match = re.search(score_pattern, review_content)

        if score_match:
            score = int(score_match.group(1))
        else:
            score = int(aspect["max_score"] * 0.82)

        score = min(aspect["max_score"], max(0, score))
        total_score += score
        dimensions.append({
            "name": aspect["name"],
            "score": score,
            "max_score": aspect["max_score"]
        })

    max_score = sum(a["max_score"] for a in config["aspects"])
    percentage = round(total_score / max_score * 100, 1)

    if percentage >= 90:
        grade = "优秀"
    elif percentage >= 80:
        grade = "良好"
    elif percentage >= 70:
        grade = "中等"
    elif percentage >= 60:
        grade = "及格"
    else:
        grade = "不及格"

    return PaperReviewResponse(
        paper_type=config["name"],
        total_score=total_score,
        grade=grade,
        review_report=review_content,
        dimensions=dimensions
    )


@router.post("/review/upload")
async def upload_and_review(
    file: UploadFile = File(...),
    paper_type: str = Form("math_modeling")
):
    """上传文件并评审"""
    allowed_types = [".pdf", ".docx", ".txt", ".md"]
    file_ext = Path(file.filename).suffix.lower()

    if file_ext not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {file_ext}。支持: {', '.join(allowed_types)}"
        )

    # 限制文件大小 (20MB)
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件过大，最大支持 20MB")

    temp_dir = Path(r"D:\MathematicModeling\workshop_output") / "temp_uploads"
    temp_dir.mkdir(exist_ok=True)
    temp_file = temp_dir / file.filename

    try:
        with open(temp_file, "wb") as f:
            f.write(content)

        if file_ext == ".docx":
            text_content = extract_text_from_docx(str(temp_file))
        elif file_ext in (".txt", ".md"):
            text_content = content.decode("utf-8", errors="ignore")
        elif file_ext == ".pdf":
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(str(temp_file))
                text_parts = []
                for page_num, page in enumerate(doc):
                    # 尝试提取文本
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                    else:
                        # 如果文本提取失败（可能是图片PDF），尝试OCR
                        try:
                            pix = page.get_pixmap(dpi=200)
                            img_data = pix.tobytes("png")
                            # 尝试用pytesseract OCR
                            try:
                                import pytesseract
                                from PIL import Image
                                import io
                                img = Image.open(io.BytesIO(img_data))
                                ocr_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                                if ocr_text.strip():
                                    text_parts.append(f"[OCR第{page_num+1}页]\n" + ocr_text)
                            except ImportError:
                                # pytesseract不可用，跳过OCR
                                text_parts.append(f"[第{page_num+1}页为图片，无法提取文本]")
                        except Exception:
                            text_parts.append(f"[第{page_num+1}页解析失败]")
                text_content = "\n".join(text_parts)
                doc.close()
            except ImportError:
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(str(temp_file))
                    text_content = "\n".join([page.extract_text() or "" for page in reader.pages])
                except ImportError:
                    raise HTTPException(
                        status_code=500,
                        detail="PDF 解析需要安装 PyMuPDF 或 PyPDF2"
                    )
        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件格式: {file_ext}")

        if not text_content.strip():
            raise HTTPException(status_code=400, detail="文件内容为空或无法解析")

        review_request = PaperReviewRequest(content=text_content, paper_type=paper_type)
        return await review_paper(review_request)

    finally:
        if temp_file.exists():
            temp_file.unlink()


@router.get("/review/configs")
async def get_review_configs():
    """获取评审配置"""
    return REVIEW_CONFIGS


@router.post("/review/export")
async def export_review_report(request: PaperReviewRequest):
    """导出评审报告为Word文档"""
    review_response = await review_paper(request)

    doc = Document()
    title = doc.add_heading('论文评审报告', 0)
    title.alignment = 1

    doc.add_paragraph(f'论文类型: {review_response.paper_type}')
    doc.add_paragraph(f'总分: {review_response.total_score} 分')
    doc.add_paragraph(f'等级: {review_response.grade}')
    doc.add_paragraph('')

    doc.add_heading('各维度评分', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = '维度'
    header_cells[1].text = '得分'
    header_cells[2].text = '满分'

    for dim in review_response.dimensions:
        row_cells = table.add_row().cells
        row_cells[0].text = dim['name']
        row_cells[1].text = str(dim['score'])
        row_cells[2].text = str(dim['max_score'])

    doc.add_paragraph('')
    doc.add_heading('评审详情', level=1)
    doc.add_paragraph(review_response.review_report)

    temp_dir = Path("temp_exports")
    temp_dir.mkdir(exist_ok=True)
    output_file = temp_dir / f"review_report_{request.paper_type}.docx"
    doc.save(str(output_file))

    return FileResponse(
        str(output_file),
        filename=f"论文评审报告_{request.paper_type}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
